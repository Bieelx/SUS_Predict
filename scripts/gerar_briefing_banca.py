from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "entregas"
DOCX = OUT / "SusPredict_Briefing_Completo_para_Banca.docx"
MD = OUT / "SusPredict_Briefing_Completo_para_Banca.md"

NAVY = "173F3A"
TEAL = "2F7469"
MINT = "DDEDE8"
INK = "202925"
MUTED = "5F6B66"
LIGHT = "F3F7F5"
AMBER = "9A6500"
RED = "9B2C2C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_font(run, size=None, bold=None, color=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, before, after, color in (
        ("Heading 1", 17, 16, 8, NAVY),
        ("Heading 2", 13.5, 12, 6, TEAL),
        ("Heading 3", 11.5, 8, 4, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10
    if "Lead" not in styles:
        style = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(12.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing = 1.15
    if "Callout" not in styles:
        style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.right_indent = Inches(0.18)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(10)
        style.paragraph_format.line_spacing = 1.10
        ppr = style.element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), MINT)
        ppr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("SusPredict | Briefing para banca   ")
    set_font(run, 8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SUSPREDICT")
    set_font(r, 12, bold=True, color=TEAL, name="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Briefing completo\npara apresentação da banca")
    set_font(r, 28, bold=True, color=NAVY, name="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("O que o projeto resolve, o que já demonstra, seus limites e o roadmap realista")
    set_font(r, 13, color=MUTED)
    p = doc.add_paragraph(style="Callout")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mensagem central: transformar dados públicos de saúde em antecedência para decisões operacionais, com transparência sobre evidências, simulações e limitações.")
    set_font(r, 12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Material-base para um agente de IA montar slides\nFIAP TCC 2025/2026 | Versão consolidada em 01/09/2026")
    set_font(r, 9.5, color=MUTED)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_p(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Situação", "O que significa", "Como apresentar"]
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_font(r, 9.5, bold=True, color=WHITE)
    mark_header_row(table.rows[0])
    for status, meaning, presentation in rows:
        cells = table.add_row().cells
        for i, value in enumerate((status, meaning, presentation)):
            p = cells[i].paragraphs[0]
            r = p.add_run(value)
            set_font(r, 9.2, bold=(i == 0), color=INK)
            if i == 0:
                set_cell_shading(cells[i], LIGHT)
    set_table_geometry(table, [1700, 3300, 4360])
    return table


def add_sources_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, value in enumerate(("Fonte", "Uso no projeto / na apresentação")):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        r = cell.paragraphs[0].add_run(value)
        set_font(r, 9.5, bold=True, color=WHITE)
    mark_header_row(table.rows[0])
    for source, use in rows:
        cells = table.add_row().cells
        for i, value in enumerate((source, use)):
            r = cells[i].paragraphs[0].add_run(value)
            set_font(r, 8.8, color=INK)
    set_table_geometry(table, [5000, 4360])


def build_markdown():
    text = """# SusPredict - briefing completo para apresentacao da banca

> Material-base para um agente de IA montar slides. Versao consolidada em 01/09/2026.

## 1. Resumo executivo

O SusPredict e uma plataforma academica de apoio a decisao para gestores de saude publica. Sua proposta e transformar dados epidemiologicos, hospitalares e sinais de aquisicao em uma leitura clara de risco, prioridade e proxima acao. Em vez de apenas exibir graficos, o produto procura responder: **o que esta acontecendo, por que importa, onde agir e qual evidencia sustenta essa decisao?**

O recorte mais forte para a banca e dengue e planejamento de insumos. A plataforma organiza dados do SINAN, SIH e tabelas curadas de compras publicas; apresenta visoes executivas e analiticas; centraliza alertas; oferece uma demo historica; permite conversar com a Clara; e conduz um alerta de insumo ate um rascunho de ETP com revisao humana.

**Frase de posicionamento:** SusPredict transforma sinais dispersos do SUS em antecedencia para decisoes de saude e compras publicas.

**Mensagem de honestidade:** a solucao apoia a decisao; nao substitui o gestor, o profissional de saude, a validacao juridica ou os sistemas oficiais.

## 2. Problema e oportunidade

Secretarias municipais precisam combinar informacoes que normalmente vivem separadas: notificacoes epidemiologicas, internacoes, aquisicoes, estoque local e prazos administrativos. Sem essa conexao, a leitura tende a ser reativa: percebe-se a pressao quando a demanda ja cresceu ou quando o insumo ja entrou em situacao critica.

O SusPredict procura reduzir esse intervalo entre o sinal e a acao. Seu diferencial nao e “ter IA” nem “ter um dashboard”; e organizar evidencias para que uma gestora nao especialista consiga compreender rapidamente a situacao, investigar a origem e iniciar uma resposta documentada.

### Publico principal

- Gestores e tecnicos de secretarias municipais e estaduais de saude.
- Epidemiologistas e equipes de vigilancia.
- Coordenadores de farmacia e profissionais envolvidos em aquisicoes.
- Pesquisadores e estudantes que precisam analisar e comunicar dados publicos de saude.

## 3. Proposta de valor

O produto une quatro capacidades em uma mesma jornada:

1. **Observar:** reunir indicadores epidemiologicos e hospitalares com fonte e competencia visiveis.
2. **Interpretar:** destacar tendencia, risco e limitacoes, evitando que o usuario precise interpretar sozinho varias tabelas.
3. **Priorizar:** transformar sinais em uma fila de alertas e itens que exigem atencao.
4. **Agir com revisao:** levar um alerta de insumo a um rascunho de ETP, sempre com confirmacao humana.

## 4. Jornada principal demonstravel

1. O usuario entra no ambiente autenticado e escolhe o municipio.
2. A Visao Geral apresenta a sintese do territorio, com fonte, competencia e comparativos.
3. O usuario aprofunda a analise em Epidemiologia, Internacoes, Vacinacao, Insumos ou Alertas.
4. A Central de Alertas prioriza sinais de risco de aquisicao por insumo.
5. A Clara explica o contexto usando ferramentas de consulta e aponta a tela relacionada.
6. Quando existe uma acao de escrita, o sistema exige confirmacao explicita.
7. O Gerador de ETP traz dados de origem, pede informacoes da secretaria e exige revisao do texto.
8. O documento pode ser baixado como PDF e aparece no historico da sessao.

## 5. O que esta implementado hoje

### Experiencia web

- Aplicacao responsiva com login, navegacao por tarefas, estados de carregamento e mensagens honestas quando a fonte esta indisponivel.
- Seletor de municipio compartilhado entre as telas municipais.
- Separacao visual entre ambiente real e demonstracao historica.
- Visao Geral, Alertas, Insumos, Documentos, Epidemiologia, Internacoes, Vacinacao, Configuracoes e Perfil.

### Dados e analises

- Endpoints autenticados para Visao Geral, Epidemiologia, Internacoes, Vacinacao e risco de aquisicao/ruptura.
- Visao Geral baseada em tabelas curadas, com KPIs, serie temporal, risco agregado, evolucao de casos, mapa/ranking regional, categorias de suprimento e alertas recentes.
- Epidemiologia municipal baseada em SINAN, com casos, incidencia, hospitalizacao, obito, perfil demografico, sazonalidade e previsao de tres meses quando a serie permite.
- Internacoes SIH apresentadas por estabelecimento/CNES e agregado estadual, coerentes com a granularidade real da fonte.
- Vacinacao contra dengue com ressalvas explicitas: serie curta, associacao nao causal e cruzamento hospitalar estadual.
- Insumos e Alertas apresentados como **risco de aquisicao**, nao como estoque fisico ou dias reais de cobertura quando esses dados locais nao existem.

### Clara

- Assistente conversacional integrado ao painel, com historico, contexto de tela e referencias de navegacao.
- Rotas deterministicas para perguntas operacionais de alta confianca e uso de modelo local/alternativo para linguagem natural.
- Ferramentas para consultar epidemiologia, estoque cadastrado, alertas e outros dados permitidos.
- Confirmacao obrigatoria antes de ferramentas que alteram estado.
- Pareamento com Telegram, continuidade de historico e suporte a audio com transcricao local quando o runtime e o modelo estao configurados.

### Demo historica e ETP

- Replay historico de dengue 2024, separado do modo real.
- Casos historicos reais; estoque, precos e economia sao um cenario demonstrativo ficticio e identificado como tal.
- Linha do tempo que evidencia a passagem de sinal epidemiologico para risco operacional.
- Fluxo de ETP em quatro etapas: dados do sistema, dados da secretaria, revisao do texto e geracao.
- Revisao humana obrigatoria antes de finalizar e baixar o PDF.

### Validacao automatizada atual

- 130 testes Python aprovados em 01/09/2026.
- Build de producao do frontend aprovado em 01/09/2026.
- Isso comprova contratos e integridade automatizada; nao equivale, por si so, a validacao clinica, juridica, de seguranca em producao ou teste visual em todos os dispositivos.

## 6. O que o SusPredict nao faz

- Nao monitora estoque fisico em tempo real sem uma fonte local de estoque e consumo.
- Nao transforma compras publicas em saldo disponivel, consumo real ou dias de cobertura.
- Nao mede ocupacao hospitalar em tempo real; o SIH registra internacoes financiadas pelo SUS e o CNES descreve capacidade/cadastro, nao ocupacao instantanea.
- Nao corrige subnotificacao nem garante que todo registro de origem esteja correto.
- Nao comprova causalidade entre vacinacao e reducao de casos ou internacoes.
- Nao diagnostica pacientes, nao prescreve condutas clinicas e nao substitui protocolos de saude.
- Nao executa automaticamente compras, licitacoes ou outras decisoes administrativas.
- Nao produz um ETP juridicamente aprovado; gera um rascunho de apoio que exige revisao tecnica e juridica.
- Nao deve apresentar dados de demo como se fossem dados operacionais presentes.
- Nao oferece, nesta fase, maturidade completa de produto multi-tenant para uso municipal em larga escala.

## 7. O que ainda esta WIP ou depende de validacao

### Produto e dados

- Integracao real com estoque e consumo local de uma secretaria, incluindo unidades, dispensacao, pedidos em transito e prazo de compra.
- Validacao clinica e operacional da relacao entre crescimento de casos e consumo de cada insumo.
- Rotina completa de importacao CSV/XLSX com mapeamento, validacao, deduplicacao, previa e desfazer.
- Validacao de campo com municipios e acompanhamento de resultados reais.
- Avaliacao mais ampla da qualidade e atualidade de cada tabela curada.

### Funcionalidades

- Persistencia completa dos ETPs e rascunhos no backend; parte do historico atual vive na sessao do frontend.
- Workflow real de aprovacao, assinatura, versionamento e auditoria documental.
- Superlotacao/ocupacao hospitalar somente quando houver fonte verificavel; a tela operacional sem fonte foi retirada.
- Notificacoes externas completas e confiaveis, com operacao permanente dos webhooks e canais.
- Integracoes futuras com sistemas municipais, RNDS e outros canais dependem de governanca, seguranca e contratos.

### Clara e infraestrutura

- Avaliacao sistematica das respostas da Clara contra um conjunto de perguntas reais.
- Validacao ao vivo e recorrente do modelo local Ollama/Qwen no ambiente Ubuntu.
- Monitoramento, observabilidade, politica de retencao e controles de custo em escala.
- Isolamento multi-tenant forte: o municipio autorizado deve vir da sessao e de politicas de acesso, nao apenas do parametro enviado pelo cliente.
- Operacao permanente do Telegram exige webhook HTTPS ativo, segredos corretos e modelo de transcricao instalado.

### Producao, seguranca e conformidade

- Piloto supervisionado, avaliacao de seguranca, LGPD, ameacas e continuidade operacional.
- Auditoria formal de acessibilidade com tecnologias assistivas reais.
- Validacao juridica do conteudo e do fluxo de ETP.
- Definicao de responsabilidade, governanca dos dados, SLAs e suporte.

## 8. Leitura correta dos dados

- **SINAN:** notificacoes e investigacoes de agravos; reflete o que foi registrado e pode sofrer atraso ou subnotificacao.
- **SIH/SUS:** autorizacoes de internacao hospitalar financiadas pelo SUS; nao e censo de ocupacao em tempo real.
- **Compras publicas curadas:** evidenciam aquisicao e fornecedores; nao equivalem a estoque atual na unidade.
- **Estoque local:** quando cadastrado, permite calculos de saldo e cobertura, mas sua confiabilidade depende de atualizacao, unidade e consumo validos.
- **Demo historica:** usa casos historicos reais e premissas ficticias claramente rotuladas para estoque, precos e economia.

## 9. Diferenciais para enfatizar na banca

- Orientacao a decisao, nao apenas visualizacao.
- Conexao entre saude publica, abastecimento e processo administrativo.
- Transparencia entre dado real, simulacao, indisponibilidade e limitacao.
- Clara como interface de explicacao e convergencia entre web e Telegram, nao como fonte magica dos numeros.
- Confirmacao humana antes de qualquer acao com impacto.
- ETP como desfecho pratico da jornada, mantendo revisao obrigatoria.

## 10. Estrutura sugerida para os slides

1. **Titulo e tese:** “Do dado disperso a decisao antecipada”.
2. **Problema:** gestores reagem tarde porque dados e processos estao separados.
3. **Publico:** a gestora municipal que precisa decidir sem uma equipe grande de dados.
4. **Solucao:** uma jornada visual simples - observar, interpretar, priorizar e agir.
5. **Como funciona:** fluxo em linguagem de negocio, sem arquitetura tecnica.
6. **Demonstracao:** replay historico de dengue, do primeiro sinal ao momento de abrir o ETP.
7. **Clara:** explicar dados e manter continuidade entre canais, sempre com guardrails.
8. **Confianca:** fontes, competencia, limites, revisao humana e separacao real/demo.
9. **O que ja existe:** produto navegavel, dados operacionais, alertas, Clara, Telegram e ETP.
10. **Limites e WIP:** estoque local, validacao clinica/juridica, piloto e escala.
11. **Roadmap:** piloto municipal, integracoes e amadurecimento de governanca.
12. **Fechamento:** “Antecipar para planejar, em vez de reagir quando a crise ja chegou”.

### Regras para o agente que criar os slides

- Priorizar historia, problema, impacto e jornada; tecnologia entra apenas como prova de viabilidade.
- Usar uma ideia principal por slide e pouco texto.
- Nao abrir a apresentacao com stack, arquitetura, endpoints, banco ou nomes de modelos.
- Nao apresentar numeros de demo como resultados reais do produto.
- Nao dizer “tempo real” para DATASUS, SIH, CNES ou compras publicas.
- Nao afirmar que a plataforma evita uma compra, reduz custo ou preve ruptura real sem piloto validado.
- Identificar sempre: implementado, demo, validado localmente, dependente de infraestrutura ou roadmap.
- Preferir capturas da jornada e diagramas simples a tabelas densas.

## 11. Roteiro recomendado da demonstracao

**Frase obrigatoria:** “Os casos de dengue sao historicos reais; estoque, precos e economia sao um cenario demonstrativo ficticio.”

1. Abrir a demo historica em janeiro de 2024.
2. Avancar para fevereiro e mostrar a aceleracao epidemiologica.
3. Avancar para marco e mostrar que a janela operacional esta diminuindo.
4. Parar em abril, destacar o item critico, o mes estimado de ruptura e a acao sugerida.
5. Abrir o Gerador de ETP, mostrar os dados de origem e a revisao humana obrigatoria.
6. Gerar o documento e abrir o historico.

O objetivo nao e explicar cada grafico. E provar que o sistema liga sinal, prioridade, evidencia e acao.

## 12. Perguntas provaveis da banca

**De onde vem o dado?** SINAN para notificacoes; SIH para internacoes financiadas pelo SUS; tabelas curadas de compras para sinais de aquisicao; e dados locais quando cadastrados. Cada tela deve mostrar fonte e competencia.

**A previsao garante que vai faltar medicamento?** Nao. Sem estoque e consumo locais atualizados, o sistema apresenta risco de aquisicao, nao ruptura fisica. A previsao e apoio a decisao, com incerteza e revisao humana.

**Qual e o papel da IA?** A Clara explica, resume, navega e prepara textos. Os numeros operacionais devem vir das fontes e dos calculos; a IA nao deve inventar indicadores nem executar acoes sem confirmacao.

**O ETP ja pode ser usado oficialmente?** Ele e um rascunho contextual para acelerar o trabalho. Precisa de revisao tecnica, juridica e adequacao ao processo do orgao.

**O sistema esta pronto para producao?** Nao. E um MVP academico funcional e demonstravel. Producao exige piloto, dados locais, isolamento entre clientes, seguranca, governanca, observabilidade e validacoes clinica e juridica.

**Qual o principal proximo passo?** Pilotar com uma secretaria que forneca estoque e consumo reais, validar os coeficientes e medir se a antecedencia gerada melhora o planejamento.

## 13. Roadmap recomendado

- **Curto prazo:** consolidar a demo da banca, corrigir integracoes operacionais e documentar cada fonte.
- **Piloto:** integrar estoque/consumo real de um municipio e validar alertas com usuarios responsaveis.
- **Produto:** persistir documentos e workflows, fortalecer isolamento de dados e operacao dos canais.
- **Escala:** integrar sistemas municipais e expandir territorios somente depois da validacao do piloto.

## 14. Fontes

### Fontes oficiais externas

1. DATASUS - Informacoes de Saude (TABNET): https://datasus.saude.gov.br/informacoes-de-saude-tabnet/
2. Portal de Dados Abertos do SUS - SINAN/Dengue: https://dadosabertos.saude.gov.br/dataset/arboviroses-dengue
3. Portal SINAN - descricao do sistema: https://www.portalsinan.saude.gov.br/
4. DATASUS - Sistema de Informacoes Hospitalares (SIH/SUS): https://siab.datasus.gov.br/DATASUS/index.php?area=060502
5. Ministerio da Saude - Transparencia dos dados SIH/SIA: https://www.gov.br/saude/pt-br/acesso-a-informacao/sic/dados-em-transparencia-ativa/saes
6. Lei 14.133/2021, especialmente art. 18 sobre fase preparatoria e ETP: https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm
7. Ministerio da Saude - Observatorio de Arboviroses: https://www.gov.br/saude/pt-br/composicao/svsa/cnie/observatorio-de-arboviroses

### Fontes internas do projeto

- `PRODUCT.md` - proposito, usuarios e principios do produto.
- `docs/01-visao-geral.md` e `docs/02-produto.md` - problema, persona, proposta e roadmap original.
- `docs/04-qualidade-dados.md` - limites e tratamentos das bases.
- `.md/AUDITORIA_PRODUTO_UX.md` - auditoria de escopo, confianca, acessibilidade e lacunas.
- `docs/06-agente-clara.md` e `docs/documentacao/API_TELEGRAM.md` - Clara, ferramentas, confirmacao e convergencia de canais.
- `docs/superpowers/demo-roteiro-2min-crise-historica-dengue.md` - roteiro da demo historica.
- Codigo atual em `api/` e `frontend/src/` - fonte de verdade do que esta implementado nesta versao.
"""
    OUT.mkdir(parents=True, exist_ok=True)
    MD.write_text(text, encoding="utf-8")
    return text


def build_docx(md_text):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    configure_styles(doc)
    add_page_number(section.footer.paragraphs[0])
    add_cover(doc)

    add_heading(doc, "Como usar este documento", 1)
    add_p(doc, "Este material foi escrito para alimentar a criação de slides. Ele privilegia a narrativa de produto e separa rigorosamente capacidade implementada, demonstração, limitação e trabalho futuro.", style="Lead")
    add_status_table(doc, [
        ("Implementado", "Existe no código atual e possui validação automatizada proporcional.", "Pode ser mostrado como capacidade do MVP."),
        ("Demo", "Funciona em cenário controlado; algumas premissas são fictícias e rotuladas.", "Apresentar como prova narrativa, não resultado operacional real."),
        ("WIP", "Existe parcialmente ou depende de dados, runtime ou integração externa.", "Mostrar como próximo passo e não como entrega concluída."),
        ("Fora de escopo", "Não é promessa desta versão.", "Usar para demonstrar maturidade e controle de escopo."),
    ])

    sections = [
        ("1. Resumo executivo", [
            ("lead", "O SusPredict é uma plataforma acadêmica de apoio à decisão para gestores de saúde pública. Sua proposta é transformar dados epidemiológicos, hospitalares e sinais de aquisição em uma leitura clara de risco, prioridade e próxima ação."),
            ("p", "Em vez de apenas exibir gráficos, o produto procura responder: o que está acontecendo, por que importa, onde agir e qual evidência sustenta essa decisão? O recorte mais forte para a banca é dengue e planejamento de insumos."),
            ("callout", "Frase de posicionamento: SusPredict transforma sinais dispersos do SUS em antecedência para decisões de saúde e compras públicas."),
            ("callout", "Limite central: a solução apoia a decisão; não substitui o gestor, o profissional de saúde, a validação jurídica ou os sistemas oficiais."),
        ]),
        ("2. Problema e oportunidade", [
            ("p", "Secretarias municipais precisam combinar informações que normalmente vivem separadas: notificações epidemiológicas, internações, aquisições, estoque local e prazos administrativos. Sem essa conexão, a leitura tende a ser reativa: percebe-se a pressão quando a demanda já cresceu ou quando o insumo já entrou em situação crítica."),
            ("p", "O diferencial não é “ter IA” nem “ter um dashboard”. É organizar evidências para que uma gestora não especialista compreenda rapidamente a situação, investigue a origem e inicie uma resposta documentada."),
            ("h2", "Público principal"),
            ("bullets", ["Gestores e técnicos de secretarias municipais e estaduais de saúde.", "Epidemiologistas e equipes de vigilância.", "Coordenadores de farmácia e profissionais envolvidos em aquisições.", "Pesquisadores e estudantes que analisam e comunicam dados públicos de saúde."]),
        ]),
        ("3. Proposta de valor", [
            ("numbered", ["Observar: reunir indicadores com fonte e competência visíveis.", "Interpretar: destacar tendência, risco e limitações em linguagem acessível.", "Priorizar: transformar sinais em uma fila de alertas e itens que exigem atenção.", "Agir com revisão: levar um alerta a um rascunho de ETP, sempre com confirmação humana."]),
            ("h2", "Diferenciais para enfatizar"),
            ("bullets", ["Orientação à decisão, não apenas visualização.", "Conexão entre saúde pública, abastecimento e processo administrativo.", "Transparência entre dado real, demo, indisponibilidade e limitação.", "Clara como interface de explicação e convergência, não como fonte mágica dos números.", "Confirmação humana antes de ações com impacto.", "ETP como desfecho prático da jornada, mantendo revisão obrigatória."]),
        ]),
        ("4. Jornada principal demonstrável", [
            ("numbered", ["Entrar no ambiente autenticado e escolher o município.", "Ler a síntese territorial na Visão Geral, com fonte e competência.", "Aprofundar a análise em Epidemiologia, Internações, Vacinação, Insumos ou Alertas.", "Priorizar sinais de risco de aquisição na Central de Alertas.", "Pedir à Clara que explique o contexto e indique a tela relacionada.", "Confirmar explicitamente qualquer ação que altere estado.", "Revisar dados e texto no Gerador de ETP.", "Baixar o PDF e consultar o histórico da sessão."]),
        ]),
        ("5. O que está implementado hoje", [
            ("h2", "Experiência web"),
            ("bullets", ["Aplicação responsiva com login, navegação por tarefas, skeletons e mensagens honestas de indisponibilidade.", "Seletor de município compartilhado entre telas municipais.", "Separação visual entre ambiente real e demonstração histórica.", "Telas de Visão Geral, Alertas, Insumos, Documentos, Epidemiologia, Internações, Vacinação, Configurações e Perfil."]),
            ("h2", "Dados e análises"),
            ("bullets", ["Contratos autenticados para as principais telas operacionais.", "Visão Geral com KPIs, série, risco agregado, evolução de casos, recorte regional, categorias e alertas.", "Epidemiologia municipal com SINAN, perfil demográfico, sazonalidade e previsão quando a série permite.", "Internações SIH por estabelecimento/CNES e agregado estadual, respeitando a granularidade da fonte.", "Vacinação com ressalvas explícitas sobre série curta, recortes e ausência de causalidade.", "Insumos e Alertas tratados como risco de aquisição quando não existe estoque local verificável."]),
            ("h2", "Clara e convergência"),
            ("bullets", ["Painel conversacional com histórico, contexto de tela e referências de navegação.", "Rotas determinísticas para perguntas operacionais de alta confiança e modelo de linguagem para formulação da resposta.", "Ferramentas de consulta com artefatos estruturados baseados no retorno real.", "Confirmação obrigatória antes de ferramentas que alteram estado.", "Pareamento com Telegram e suporte a áudio com transcrição local quando a infraestrutura está configurada."]),
            ("h2", "Demo histórica e ETP"),
            ("bullets", ["Replay de dengue 2024 separado do modo real.", "Casos históricos reais; estoque, preços e economia são premissas fictícias e rotuladas.", "Fluxo de ETP em quatro etapas, com texto editável e revisão humana obrigatória.", "Download em PDF e histórico da sessão."]),
            ("callout", "Validação automatizada em 01/09/2026: 130 testes Python aprovados e build de produção do frontend aprovado. Isso não equivale a validação clínica, jurídica, visual completa ou prontidão para produção."),
        ]),
        ("6. O que o SusPredict não faz", [
            ("bullets", ["Não monitora estoque físico em tempo real sem fonte local de estoque e consumo.", "Não transforma compras públicas em saldo, consumo real ou dias de cobertura.", "Não mede ocupação hospitalar em tempo real.", "Não corrige subnotificação nem garante que todo registro de origem esteja correto.", "Não comprova causalidade entre vacinação e desfechos.", "Não diagnostica pacientes nem prescreve condutas clínicas.", "Não executa automaticamente compras ou licitações.", "Não produz um ETP juridicamente aprovado; gera um rascunho para revisão.", "Não deve apresentar dados de demo como situação operacional presente.", "Não possui ainda maturidade multi-tenant para escala produtiva."]),
        ]),
        ("7. O que ainda está WIP", [
            ("h2", "Produto e dados"),
            ("bullets", ["Integração real com estoque, consumo, pedidos em trânsito e prazo de compra de uma secretaria.", "Validação clínica e operacional da relação entre casos e consumo de cada insumo.", "Importação CSV/XLSX com mapeamento, validação, deduplicação, prévia e desfazer.", "Piloto com município e medição de resultados reais."]),
            ("h2", "Funcionalidades"),
            ("bullets", ["Persistência completa de ETPs e rascunhos no backend.", "Workflow de aprovação, assinatura, versionamento e auditoria documental.", "Ocupação/superlotação apenas quando existir fonte verificável.", "Notificações externas com operação permanente dos webhooks e canais."]),
            ("h2", "Clara, infraestrutura e produção"),
            ("bullets", ["Avaliação sistemática da Clara com perguntas reais.", "Validação ao vivo recorrente do runtime Ollama/Qwen no Ubuntu.", "Observabilidade, retenção, monitoramento e controles de custo.", "Isolamento multi-tenant forte, derivado da sessão e de políticas de acesso.", "Piloto supervisionado, segurança, LGPD, acessibilidade formal e validação jurídica."]),
        ]),
        ("8. Como ler corretamente os dados", [
            ("bullets", ["SINAN: notificações e investigações; pode haver atraso e subnotificação.", "SIH/SUS: autorizações de internação financiadas pelo SUS; não é ocupação em tempo real.", "Compras públicas curadas: evidenciam aquisição e fornecedores; não equivalem a estoque atual.", "Estoque local: permite saldo e cobertura quando atualizado e com unidades/consumo válidos.", "Demo histórica: casos reais e premissas fictícias explicitamente rotuladas para estoque, preços e economia."]),
        ]),
        ("9. Estrutura sugerida para os slides", [
            ("numbered", ["Título e tese: do dado disperso à decisão antecipada.", "Problema: dados e processos separados levam a reações tardias.", "Público: a gestora que precisa decidir sem uma grande equipe de dados.", "Solução: observar, interpretar, priorizar e agir.", "Como funciona: fluxo em linguagem de negócio.", "Demonstração: replay de dengue, do sinal ao ETP.", "Clara: explicação e continuidade entre canais com guardrails.", "Confiança: fontes, competência, limites e revisão humana.", "O que já existe: produto navegável e capacidades demonstráveis.", "Limites e WIP: dados locais, validações, piloto e escala.", "Roadmap: do MVP acadêmico ao piloto municipal.", "Fechamento: antecipar para planejar, em vez de reagir quando a crise chegou."]),
            ("h2", "Regras para o agente de slides"),
            ("bullets", ["Uma ideia principal por slide e pouco texto.", "Priorizar história, problema, impacto e jornada; tecnologia entra como prova de viabilidade.", "Não abrir com stack, endpoints, banco ou nomes de modelos.", "Não apresentar valores da demo como resultados reais.", "Não dizer “tempo real” para DATASUS, SIH, CNES ou compras públicas.", "Não prometer redução de custo ou previsão de ruptura real sem piloto validado.", "Identificar sempre: implementado, demo, validação local, dependência externa ou roadmap.", "Preferir capturas da jornada e diagramas simples a tabelas densas."]),
        ]),
        ("10. Roteiro recomendado da demonstração", [
            ("callout", "Frase obrigatória: Os casos de dengue são históricos reais; estoque, preços e economia são um cenário demonstrativo fictício."),
            ("numbered", ["Abrir a demo histórica em janeiro de 2024.", "Avançar para fevereiro e mostrar a aceleração epidemiológica.", "Avançar para março e mostrar a janela operacional diminuindo.", "Parar em abril e destacar item crítico, ruptura estimada e ação sugerida.", "Abrir o Gerador de ETP e mostrar origem e revisão obrigatória.", "Gerar o documento e abrir o histórico."]),
            ("p", "O objetivo não é explicar cada gráfico. É provar que o sistema liga sinal, prioridade, evidência e ação."),
        ]),
        ("11. Perguntas prováveis da banca", [
            ("h2", "De onde vem o dado?"), ("p", "SINAN para notificações; SIH para internações financiadas pelo SUS; tabelas curadas de compras para sinais de aquisição; e dados locais quando cadastrados. Cada tela deve mostrar fonte e competência."),
            ("h2", "A previsão garante que vai faltar medicamento?"), ("p", "Não. Sem estoque e consumo locais atualizados, o sistema apresenta risco de aquisição, não ruptura física. A previsão é apoio à decisão e exige revisão humana."),
            ("h2", "Qual é o papel da IA?"), ("p", "A Clara explica, resume, navega e prepara textos. Os números devem vir das fontes e dos cálculos; a IA não executa ações sem confirmação."),
            ("h2", "O ETP já pode ser usado oficialmente?"), ("p", "É um rascunho contextual para acelerar o trabalho. Precisa de revisão técnica, jurídica e adequação ao processo do órgão."),
            ("h2", "O sistema está pronto para produção?"), ("p", "Não. É um MVP acadêmico funcional e demonstrável. Produção exige piloto, dados locais, segurança, isolamento entre clientes, governança e validações clínica e jurídica."),
            ("h2", "Qual o principal próximo passo?"), ("p", "Pilotar com uma secretaria que forneça estoque e consumo reais, validar os coeficientes e medir se a antecedência melhora o planejamento."),
        ]),
        ("12. Roadmap recomendado", [
            ("bullets", ["Curto prazo: consolidar a demo, corrigir integrações e documentar cada fonte.", "Piloto: integrar estoque/consumo real de um município e validar alertas com responsáveis.", "Produto: persistir documentos e workflows, fortalecer isolamento e operação dos canais.", "Escala: integrar sistemas municipais e expandir territórios somente depois da validação do piloto."]),
        ]),
    ]

    for title, blocks in sections:
        add_heading(doc, title, 1)
        for kind, value in blocks:
            if kind == "h2": add_heading(doc, value, 2)
            elif kind == "p": add_p(doc, value)
            elif kind == "lead": add_p(doc, value, "Lead")
            elif kind == "callout": add_p(doc, value, "Callout")
            elif kind == "bullets": add_bullets(doc, value)
            elif kind == "numbered": add_numbered(doc, value)

    add_heading(doc, "13. Fontes", 1)
    add_p(doc, "Fontes oficiais externas", style="Lead")
    add_sources_table(doc, [
        ("DATASUS - Informações de Saúde (TABNET)\nhttps://datasus.saude.gov.br/informacoes-de-saude-tabnet/", "Disponibilidade oficial de informações epidemiológicas e assistenciais, incluindo SINAN e SIH/SUS."),
        ("Portal de Dados Abertos do SUS - SINAN/Dengue\nhttps://dadosabertos.saude.gov.br/dataset/arboviroses-dengue", "Origem oficial dos arquivos abertos de dengue e descrição da finalidade do SINAN."),
        ("Portal SINAN\nhttps://www.portalsinan.saude.gov.br/", "Descrição institucional do sistema de notificação e investigação de agravos."),
        ("DATASUS - SIH/SUS\nhttps://siab.datasus.gov.br/DATASUS/index.php?area=060502", "Finalidade da AIH e do registro de internações financiadas pelo SUS."),
        ("Ministério da Saúde - Transparência SIH/SIA\nhttps://www.gov.br/saude/pt-br/acesso-a-informacao/sic/dados-em-transparencia-ativa/saes", "Confirma a consulta pública dos dados de produção hospitalar e ambulatorial."),
        ("Lei 14.133/2021\nhttps://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm", "Base legal da fase preparatória e do estudo técnico preliminar; não substitui parecer jurídico."),
        ("Observatório de Arboviroses - Ministério da Saúde\nhttps://www.gov.br/saude/pt-br/composicao/svsa/cnie/observatorio-de-arboviroses", "Contexto oficial sobre detecção precoce, tendências e apoio à decisão em arboviroses."),
    ])
    add_heading(doc, "Fontes internas do projeto", 2)
    add_bullets(doc, ["PRODUCT.md - propósito, usuários e princípios do produto.", "docs/01-visao-geral.md e docs/02-produto.md - problema, persona, proposta e roadmap original.", "docs/04-qualidade-dados.md - limites e tratamentos das bases.", ".md/AUDITORIA_PRODUTO_UX.md - auditoria de escopo, confiança, acessibilidade e lacunas.", "docs/06-agente-clara.md e docs/documentacao/API_TELEGRAM.md - Clara, ferramentas, confirmação e canais.", "docs/superpowers/demo-roteiro-2min-crise-historica-dengue.md - roteiro da demo.", "Código atual em api/ e frontend/src/ - fonte de verdade do que está implementado nesta versão."])

    props = doc.core_properties
    props.title = "SusPredict - Briefing completo para apresentação da banca"
    props.subject = "Produto, escopo, limites, WIP e orientação para criação de slides"
    props.author = "Equipe SusPredict"
    props.keywords = "SusPredict, FIAP, banca, DATASUS, dengue, Clara, ETP"
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)


if __name__ == "__main__":
    md = build_markdown()
    build_docx(md)
    print(DOCX)
    print(MD)
